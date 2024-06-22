import io
from typing import IO
from docx import Document


class WordDocument:

    def __init__(self, document: str | IO[bytes] | None = None):
        self.document = Document(document)

    @property
    def text(self) -> str:
        out = ''
        for para in self.document.paragraphs:
            out += para.text + '\n'
        return out

    @staticmethod
    def docx(text: str) -> io.BytesIO:
        document = Document()
        for line in text.splitlines():
            document.add_paragraph(line.strip())

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

